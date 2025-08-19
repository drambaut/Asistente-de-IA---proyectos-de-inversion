from flask import Flask, render_template, request, jsonify, session, send_from_directory, url_for
from flask_cors import CORS
import os
from datetime import datetime
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import AzureOpenAI
from dotenv import load_dotenv
import json
import time

# Configurar logging
logging.basicConfig(level=logging.INFO)  # Cambiar a INFO para producción
logger = logging.getLogger(__name__)

# Cargar variables de entorno
load_dotenv()

app = Flask(__name__, static_folder='static')
CORS(app)  # Habilitar CORS para todos los endpoints

# Configurar clave secreta desde variable de entorno o usar una por defecto
app.secret_key = os.getenv('SECRET_KEY', 'idec_secret_key_change_in_production')

# Asegurarse de que el directorio de documentos existe
DOCUMENTS_DIR = os.path.join(app.static_folder, 'documents')
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

# Configurar Azure OpenAI Client
client = AzureOpenAI(
    api_key=os.getenv('OPENAI_API_KEY'),
    api_version=os.getenv('OPENAI_API_VERSION', '2024-02-15-preview'),
    azure_endpoint=os.getenv('OPENAI_API_BASE')
)

# Obtener el ID del asistente existente
ASSISTANT_ID = os.getenv('ASSISTANT_ID')

# Validar que las variables de entorno estén configuradas
if not os.getenv('OPENAI_API_KEY'):
    raise ValueError("OPENAI_API_KEY no está configurada en el archivo .env")
if not os.getenv('OPENAI_API_BASE'):
    raise ValueError("OPENAI_API_BASE no está configurada en el archivo .env")
if not ASSISTANT_ID:
    raise ValueError("ASSISTANT_ID no está configurada en el archivo .env")

# Definir el nuevo flujo de conversación MGA/IDEC/IA
conversation_flow = {
    'welcome_message': {
        'prompt': '''👋 ¡Hola! Soy tu asistente virtual para ayudarte en la formulación de proyectos de inversión relacionados con Infraestructura de Datos (IDEC) o Inteligencia Artificial (IA). Vamos a empezar paso a paso.

Te acompañaré paso a paso para estructurar tu proyecto conforme a la Metodología General Ajustada (MGA) del Departamento Nacional de Planeación, incorporando los enfoques técnicos y estratégicos de las guías de Infraestructura de datos del Estado Colombiano (IDEC) e Inteligencia artificial.

🧰 A lo largo del proceso, te haré preguntas que nos permitirán construir los elementos clave del proyecto: desde la definición del problema, identificación de causas y objetivos (árboles de problema), hasta la justificación, la población beneficiaria, y el desarrollo de objetivos, cadena de valor, indicadores, presupuesto, cronograma y demás componentes técnicos.

❓ ¿Tienes dudas generales antes de empezar?

Antes de continuar, me gustaría saber si tienes claridad sobre el proceso general de formulación de proyectos. Si lo necesitas, puedo explicarte brevemente en qué consiste cada fase del ciclo de inversión pública:

Por favor, selecciona una opción:''',
        'options': [
            'Sí, entiendo el proceso y deseo continuar',
            'No del todo, me gustaría una breve explicación',
            'Tengo dudas puntuales sobre los lineamientos del Plan Nacional de Infraestructura de Datos (PNID) o del CONPES 4144 de Inteligencia Artificial'
        ],
        'next_step': 'handle_welcome_choice'
    },
    'cycle_question': {
        'prompt': '''¿Conoces el ciclo de inversión pública y las fases que lo componen?''',
        'options': [
            'Sí, lo conozco',
            'No, me gustaría entenderlo mejor'
        ],
        'next_step': 'handle_cycle_choice'
    },
    'cycle_explanation': {
        'prompt': '''No te preocupes. El ciclo de inversión pública incluye las siguientes etapas:

• Identificación del problema u oportunidad de mejora
• Formulación de alternativas y estructuración técnica y financiera
• Evaluación y viabilidad del proyecto
• Registro en el Banco de Programas y Proyectos (BPIN)
• Implementación, seguimiento y evaluación

📘 Puedes conocer más en la Guía MGA del DNP

¿Estás listo para comenzar con la formulación de tu proyecto?''',
        'options': [
            'Sí, comenzar con la formulación'
        ],
        'next_step': 'start_formulation'
    },
    'brief_explanation': {
        'prompt': '''📋 **Breve explicación del ciclo de inversión pública:**

El ciclo de inversión pública consta de las siguientes fases principales:

1. **Pre-inversión**: Identificación, preparación y evaluación del proyecto
2. **Inversión**: Ejecución del proyecto
3. **Operación**: Funcionamiento y mantenimiento del proyecto

En la fase de **Pre-inversión** (donde nos enfocaremos), seguiremos la Metodología General Ajustada (MGA) que incluye:

• **Identificación**: Definir el problema y sus causas
• **Preparación**: Formular alternativas de solución
• **Evaluación**: Analizar viabilidad técnica, financiera y social

Para proyectos IDEC/IA, incorporaremos elementos específicos como:
- Arquitectura de datos
- Interoperabilidad
- Seguridad y privacidad
- Gobernanza de datos
- Componentes de IA

¿Estás listo para comenzar con la formulación de tu proyecto?''',
        'options': [
            'Sí, comenzar con la formulación',
            'Necesito más información sobre IDEC/IA'
        ],
        'next_step': 'handle_explanation_choice'
    },
    'idec_ia_info': {
        'prompt': '''📚 **Información sobre lineamientos IDEC e IA:**

**Plan Nacional de Infraestructura de Datos (PNID):**
- Marco estratégico para la gestión de datos del Estado
- Enfoque en interoperabilidad, calidad y seguridad de datos
- Componentes: Gobernanza, arquitectura, estándares y herramientas

**CONPES 4144 - Política Nacional de Inteligencia Artificial:**
- Estrategia nacional para adopción responsable de IA
- Pilares: Talento humano, investigación, marco ético y regulatorio
- Sectores prioritarios: Salud, educación, justicia, agricultura

**Componentes técnicos clave para tu proyecto:**
• Arquitectura de datos y sistemas
• Modelos de IA y algoritmos
• Interfaces y APIs
• Seguridad y privacidad
• Monitoreo y evaluación

¿Estás listo para iniciar la formulación considerando estos lineamientos?''',
        'options': [
            'Sí, iniciar formulación con enfoque IDEC/IA',
            'Necesito más detalles específicos'
        ],
        'next_step': 'handle_info_choice'
    },
    'entidad_nombre': {
        'prompt': '''🏛️ **Información de la Entidad Ejecutora**

Comencemos identificando la entidad que ejecutará el proyecto.

Por favor, ingrese el **nombre completo de la entidad pública** que será responsable de la ejecución del proyecto:''',
        'next_step': 'entidad_sector'
    },
    'entidad_sector': {
        'prompt': '''🏢 **Sector de la Entidad**

Por favor, especifique el **sector al que pertenece la entidad** (ejemplo: Salud, Educación, Hacienda, Tecnologías de la Información, etc.):''',
        'next_step': 'tipo_proyecto'
    },
    'tipo_proyecto': {
        'prompt': '''🎯 **Tipo de Proyecto**

Para estructurar adecuadamente tu proyecto, necesito conocer el enfoque principal.

¿Tu proyecto se enfoca principalmente en:''',
        'options': [
            'Infraestructura de Datos (IDEC) - Gestión, interoperabilidad, calidad de datos',
            'Inteligencia Artificial - Modelos, algoritmos, automatización',
            'Proyecto híbrido - Combina elementos de IDEC e IA'
        ],
        'next_step': 'problema_identificacion'
    },
    'problema_identificacion': {
        'prompt': '''🎯 **Identificación del Problema**

Ahora vamos a definir el problema que tu proyecto busca resolver. Esta es la base fundamental de toda la formulación.

**Describe detalladamente el problema o necesidad** que el proyecto pretende atender:

- ¿Cuál es la situación problemática actual?
- ¿Cómo se manifiesta este problema?
- ¿Qué evidencia tienes de que existe este problema?

(Máximo 800 caracteres)''',
        'next_step': 'causas_problema'
    },
    'causas_problema': {
        'prompt': '''🌳 **Análisis de Causas**

Para construir el árbol de problemas, identifiquemos las causas que originan la problemática.

**Lista las principales causas** que generan el problema identificado:

- Causas directas (que provocan inmediatamente el problema)
- Causas indirectas (que están detrás de las causas directas)
- Considera aspectos técnicos, organizacionales, normativos, de recursos, etc.

(Máximo 800 caracteres)''',
        'next_step': 'efectos_problema'
    },
    'efectos_problema': {
        'prompt': '''📊 **Análisis de Efectos**

Ahora identifiquemos los efectos o consecuencias que genera el problema.

**Describe los principales efectos** que produce la problemática identificada:

- Efectos directos (consecuencias inmediatas del problema)
- Efectos indirectos (consecuencias de los efectos directos)
- Impactos en la población, en la entidad, en otros sectores

(Máximo 800 caracteres)''',
        'next_step': 'poblacion_afectada'
    },
    'poblacion_afectada': {
        'prompt': '''👥 **Población Afectada**

Identifiquemos quiénes se ven impactados por esta problemática.

**Describe la población afectada** por el problema:

- ¿Quiénes son los afectados? (ciudadanos, funcionarios, empresas, etc.)
- ¿Cuántas personas aproximadamente?
- ¿En qué ubicación geográfica?
- ¿Qué características específicas tienen?

(Máximo 600 caracteres)''',
        'next_step': 'objetivo_general'
    },
    'objetivo_general': {
        'prompt': '''🎯 **Objetivo General**

Con base en el problema identificado, definamos el objetivo principal del proyecto.

**Formula el objetivo general** que el proyecto pretende alcanzar:

- Debe ser claro, medible y alcanzable
- Debe responder directamente al problema identificado
- Utiliza verbos en infinitivo (mejorar, implementar, desarrollar, etc.)

(Máximo 400 caracteres)''',
        'next_step': 'objetivos_especificos'
    },
    'objetivos_especificos': {
        'prompt': '''📋 **Objetivos Específicos**

Los objetivos específicos detallan cómo se logrará el objetivo general.

**Lista los objetivos específicos** del proyecto:

- Deben ser concretos y medibles
- Cada uno debe contribuir al logro del objetivo general
- Considera aspectos técnicos, de implementación, de capacitación, etc.

(Máximo 800 caracteres)''',
        'next_step': 'justificacion'
    },
    'justificacion': {
        'prompt': '''📝 **Justificación del Proyecto**

Expliquemos por qué es importante ejecutar este proyecto.

**Proporciona la justificación** del proyecto considerando:

- Importancia del problema para la entidad/sector
- Beneficios esperados de la solución
- Alineación con políticas públicas (PNID, CONPES 4144, etc.)
- Urgencia de la intervención

(Máximo 1000 caracteres)''',
        'next_step': 'localizacion'
    },
    'localizacion': {
        'prompt': '''📍 **Localización del Proyecto**

Definamos dónde se ejecutará el proyecto.

**Describe la localización** del proyecto:

- Cobertura geográfica (nacional, departamental, municipal, etc.)
- Ubicaciones específicas donde se implementará
- Justificación de la localización elegida

(Máximo 400 caracteres)''',
        'next_step': 'alternativas_solucion'
    },
    'alternativas_solucion': {
        'prompt': '''💡 **Alternativas de Solución**

Identifiquemos las posibles alternativas para resolver el problema.

**Describe las alternativas de solución** evaluadas:

- Alternativa 1: ¿Cuál sería?
- Alternativa 2: ¿Existe otra opción?
- ¿Por qué se selecciona una alternativa específica?
- Considera aspectos técnicos, de costos, de tiempo

(Máximo 800 caracteres)''',
        'next_step': 'componentes_tecnicos'
    },
    'componentes_tecnicos': {
        'prompt': '''⚙️ **Componentes Técnicos**

Definamos los componentes técnicos específicos del proyecto IDEC/IA.

**Describe los componentes técnicos principales**:

Para proyectos IDEC: Arquitectura de datos, APIs, estándares, seguridad
Para proyectos IA: Modelos, algoritmos, datasets, infraestructura computacional
Para proyectos híbridos: Combinación de ambos

(Máximo 800 caracteres)''',
        'next_step': 'cadena_valor'
    },
    'cadena_valor': {
        'prompt': '''🔗 **Cadena de Valor**

Describamos cómo el proyecto genera valor desde los insumos hasta los impactos.

**Define la cadena de valor** del proyecto:

- **Insumos**: Recursos necesarios (humanos, técnicos, financieros)
- **Actividades**: Qué se hará específicamente
- **Productos**: Qué se entregará
- **Resultados**: Cambios esperados en el corto plazo
- **Impactos**: Efectos de largo plazo

(Máximo 1000 caracteres)''',
        'next_step': 'indicadores'
    },
    'indicadores': {
        'prompt': '''📊 **Indicadores de Seguimiento**

Definamos cómo mediremos el éxito del proyecto.

**Propone indicadores** para medir:

- **Indicadores de producto**: ¿Qué entregas mediremos?
- **Indicadores de resultado**: ¿Qué cambios mediremos?
- **Indicadores de impacto**: ¿Qué efectos de largo plazo mediremos?

Para cada indicador incluye: nombre, unidad de medida, meta

(Máximo 800 caracteres)''',
        'next_step': 'riesgos'
    },
    'riesgos': {
        'prompt': '''⚠️ **Análisis de Riesgos**

Identifiquemos los principales riesgos que pueden afectar el proyecto.

**Lista y describe los riesgos identificados**:

- Riesgos técnicos (tecnología, implementación)
- Riesgos de recursos (presupuesto, personal)
- Riesgos normativos (cambios regulatorios)
- Riesgos organizacionales (resistencia al cambio)

Para cada riesgo indica: probabilidad, impacto y medida de mitigación

(Máximo 800 caracteres)''',
        'next_step': 'sostenibilidad'
    },
    'sostenibilidad': {
        'prompt': '''🌱 **Estrategia de Sostenibilidad**

Definamos cómo se mantendrá el proyecto en el tiempo.

**Describe la estrategia de sostenibilidad**:

- **Sostenibilidad técnica**: Mantenimiento, actualizaciones
- **Sostenibilidad financiera**: Fuentes de financiación continua
- **Sostenibilidad institucional**: Capacidades, procesos
- **Sostenibilidad ambiental**: Impacto ecológico

(Máximo 800 caracteres)''',
        'next_step': 'presupuesto_general'
    },
    'presupuesto_general': {
        'prompt': '''💰 **Presupuesto General**

Estimemos el costo total del proyecto.

**Proporciona el presupuesto general**:

- Costo total estimado del proyecto
- Fuentes de financiación (recursos propios, crédito, cooperación, etc.)
- Distribución por años (si aplica)
- Justificación del monto estimado

(Máximo 600 caracteres)''',
        'next_step': 'presupuesto_detallado'
    },
    'presupuesto_detallado': {
        'prompt': '''📋 **Presupuesto Detallado**

Desglosa el presupuesto por componentes principales.

**Detalla el presupuesto por rubros**:

- Personal (salarios, honorarios)
- Tecnología (hardware, software, licencias)
- Capacitación y formación
- Infraestructura
- Otros gastos operativos

(Máximo 800 caracteres)''',
        'next_step': 'cronograma'
    },
    'cronograma': {
        'prompt': '''📅 **Cronograma del Proyecto**

Definamos los tiempos de ejecución del proyecto.

**Describe el cronograma general**:

- Duración total del proyecto
- Fases principales y su duración
- Hitos importantes y fechas clave
- Actividades críticas que no pueden retrasarse

(Máximo 600 caracteres)''',
        'next_step': 'end'
    }
}

@app.route('/')
def index():
    # Inicializar la sesión
    session.clear()
    session['current_step'] = 'initial'
    session['responses'] = {}
    return render_template('index.html')

@app.route('/api/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        current_step = session.get('current_step', 'initial')
        
        logger.debug(f"Paso actual: {current_step}, Mensaje: {user_message}")
        
        # Paso inicial - mostrar mensaje de bienvenida
        if current_step == 'initial':
            session['current_step'] = 'welcome_message'
            return jsonify({
                'response': conversation_flow['welcome_message']['prompt'],
                'current_step': 'welcome_message',
                'options': conversation_flow['welcome_message']['options']
            })

        # Manejar las opciones de bienvenida
        if current_step == 'welcome_message':
            return handle_welcome_choice(user_message)
        
        # Manejar la pregunta sobre el ciclo de inversión
        if current_step == 'cycle_question':
            return handle_cycle_choice(user_message)
        
        # Manejar después de la explicación del ciclo
        if current_step == 'cycle_explanation':
            return handle_cycle_explanation_response(user_message)
        
        # Manejar opciones después de explicación
        if current_step == 'brief_explanation':
            return handle_explanation_choice_response(user_message)
        
        # Manejar opciones después de información IDEC/IA
        if current_step == 'idec_ia_info':
            return handle_info_choice_response(user_message)

        # Para pasos con opciones múltiples
        if current_step in conversation_flow and 'options' in conversation_flow[current_step]:
            return handle_multiple_choice(current_step, user_message)

        # Almacenar la respuesta actual para pasos de texto libre
        session['responses'] = session.get('responses', {})
        session['responses'][current_step] = user_message

        # Obtener el siguiente paso
        if current_step in conversation_flow:
            next_step = conversation_flow[current_step]['next_step']
            session['current_step'] = next_step

            # Si es el paso final, generar el documento
            if next_step == 'end':
                try:
                    # Generar el documento usando el asistente
                    document_content = generate_document_with_assistant(session['responses'])
                    filename = save_document(document_content)
                    download_url = url_for('download_file', filename=filename)
                    
                    return jsonify({
                        'response': f'¡Gracias por proporcionar toda la información! Su documento ha sido generado conforme a la metodología MGA y lineamientos IDEC/IA. <a href="{download_url}" class="download-link">📄 Descargar Documento</a>',
                        'current_step': 'end',
                        'download_url': download_url
                    })
                except Exception as e:
                    logger.error(f"Error generando documento: {str(e)}")
                    return jsonify({
                        'response': 'Lo siento, ha ocurrido un error al generar el documento. Por favor, intente nuevamente.',
                        'current_step': 'end'
                    })

            # Continuar con el siguiente paso
            response_data = {
                'response': conversation_flow[next_step]['prompt'],
                'current_step': next_step
            }
            
            # Agregar opciones si las hay
            if 'options' in conversation_flow[next_step]:
                response_data['options'] = conversation_flow[next_step]['options']
            
            return jsonify(response_data)

        else:
            return jsonify({
                'response': 'Lo siento, ha ocurrido un error en el flujo de conversación.',
                'current_step': 'end'
            })

    except Exception as e:
        logger.error(f"Error en chat: {str(e)}")
        return jsonify({
            'response': 'Lo siento, ha ocurrido un error inesperado. Por favor, intente nuevamente.',
            'current_step': 'end'
        })

def handle_welcome_choice(user_message):
    """Maneja la elección de bienvenida del usuario"""
    user_message_lower = user_message.lower()
    
    if 'entiendo' in user_message_lower and 'continuar' in user_message_lower:
        # Usuario quiere continuar directamente - hacer la pregunta sobre el ciclo
        session['current_step'] = 'cycle_question'
        return jsonify({
            'response': conversation_flow['cycle_question']['prompt'],
            'current_step': 'cycle_question',
            'options': conversation_flow['cycle_question']['options']
        })
    elif 'breve explicación' in user_message_lower or 'explicación' in user_message_lower:
        # Usuario quiere explicación
        session['current_step'] = 'brief_explanation'
        return jsonify({
            'response': conversation_flow['brief_explanation']['prompt'],
            'current_step': 'brief_explanation',
            'options': conversation_flow['brief_explanation']['options']
        })
    elif 'dudas puntuales' in user_message_lower or 'pnid' in user_message_lower or 'conpes' in user_message_lower:
        # Usuario quiere información sobre IDEC/IA
        session['current_step'] = 'idec_ia_info'
        return jsonify({
            'response': conversation_flow['idec_ia_info']['prompt'],
            'current_step': 'idec_ia_info',
            'options': conversation_flow['idec_ia_info']['options']
        })
    else:
        # Respuesta no reconocida, mostrar opciones nuevamente
        return jsonify({
            'response': 'Por favor, selecciona una de las opciones disponibles:',
            'current_step': 'welcome_message',
            'options': conversation_flow['welcome_message']['options']
        })

def handle_cycle_choice(user_message):
    """Maneja la respuesta sobre el conocimiento del ciclo de inversión"""
    user_message_lower = user_message.lower()
    
    if 'sí' in user_message_lower or 'si' in user_message_lower or 'lo conozco' in user_message_lower:
        # Usuario conoce el ciclo, continuar directamente
        session['current_step'] = 'entidad_nombre'
        return jsonify({
            'response': conversation_flow['entidad_nombre']['prompt'],
            'current_step': 'entidad_nombre'
        })
    elif 'no' in user_message_lower or 'entenderlo mejor' in user_message_lower:
        # Usuario no conoce el ciclo, mostrar explicación
        session['current_step'] = 'cycle_explanation'
        return jsonify({
            'response': conversation_flow['cycle_explanation']['prompt'],
            'current_step': 'cycle_explanation',
            'options': conversation_flow['cycle_explanation']['options']
        })
    else:
        # Respuesta no reconocida, mostrar opciones nuevamente
        return jsonify({
            'response': 'Por favor, selecciona una de las opciones disponibles:',
            'current_step': 'cycle_question',
            'options': conversation_flow['cycle_question']['options']
        })

def handle_cycle_explanation_response(user_message):
    """Maneja la respuesta después de la explicación del ciclo"""
    user_message_lower = user_message.lower()
    
    if 'sí' in user_message_lower or 'si' in user_message_lower or 'comenzar' in user_message_lower:
        # Usuario quiere comenzar la formulación
        session['current_step'] = 'entidad_nombre'
        return jsonify({
            'response': conversation_flow['entidad_nombre']['prompt'],
            'current_step': 'entidad_nombre'
        })
    else:
        # Respuesta no reconocida, mostrar opciones nuevamente
        return jsonify({
            'response': 'Por favor, selecciona una de las opciones disponibles:',
            'current_step': 'cycle_explanation',
            'options': conversation_flow['cycle_explanation']['options']
        })

def handle_explanation_choice_response(user_message):
    """Maneja la respuesta después de la explicación"""
    user_message_lower = user_message.lower()
    
    if 'comenzar' in user_message_lower or 'formulación' in user_message_lower:
        # Después de la explicación, hacer la pregunta sobre el ciclo
        session['current_step'] = 'cycle_question'
        return jsonify({
            'response': conversation_flow['cycle_question']['prompt'],
            'current_step': 'cycle_question',
            'options': conversation_flow['cycle_question']['options']
        })
    elif 'más información' in user_message_lower or 'idec' in user_message_lower:
        session['current_step'] = 'idec_ia_info'
        return jsonify({
            'response': conversation_flow['idec_ia_info']['prompt'],
            'current_step': 'idec_ia_info',
            'options': conversation_flow['idec_ia_info']['options']
        })
    else:
        return jsonify({
            'response': 'Por favor, selecciona una de las opciones disponibles:',
            'current_step': 'brief_explanation',
            'options': conversation_flow['brief_explanation']['options']
        })

def handle_info_choice_response(user_message):
    """Maneja la respuesta después de la información IDEC/IA"""
    user_message_lower = user_message.lower()
    
    if 'iniciar' in user_message_lower or 'formulación' in user_message_lower:
        # Después de la información IDEC/IA, hacer la pregunta sobre el ciclo
        session['current_step'] = 'cycle_question'
        return jsonify({
            'response': conversation_flow['cycle_question']['prompt'],
            'current_step': 'cycle_question',
            'options': conversation_flow['cycle_question']['options']
        })
    else:
        return jsonify({
            'response': 'Por favor, selecciona una de las opciones disponibles:',
            'current_step': 'idec_ia_info',
            'options': conversation_flow['idec_ia_info']['options']
        })

def handle_multiple_choice(current_step, user_message):
    """Maneja pasos con opciones múltiples"""
    # Almacenar la respuesta
    session['responses'] = session.get('responses', {})
    session['responses'][current_step] = user_message
    
    # Continuar al siguiente paso
    next_step = conversation_flow[current_step]['next_step']
    session['current_step'] = next_step
    
    response_data = {
        'response': conversation_flow[next_step]['prompt'],
        'current_step': next_step
    }
    
    # Agregar opciones si las hay
    if 'options' in conversation_flow[next_step]:
        response_data['options'] = conversation_flow[next_step]['options']
    
    return jsonify(response_data)

def generate_document_with_assistant(responses):
    """Genera el documento usando tu asistente existente de Azure OpenAI"""
    try:
        logger.info(f"Usando asistente con ID: {ASSISTANT_ID}")
        
        # Crear un thread para esta conversación
        thread = client.beta.threads.create()
        
        # Crear el mensaje con toda la información del proyecto
        message_content = f"""
        Por favor, genera un documento completo de proyecto de inversión conforme a la Metodología General Ajustada (MGA) y lineamientos IDEC/IA basado en la siguiente información:

        INFORMACIÓN DE LA ENTIDAD:
        - Entidad Ejecutora: {responses.get('entidad_nombre', '')}
        - Sector: {responses.get('entidad_sector', '')}
        - Tipo de Proyecto: {responses.get('tipo_proyecto', '')}

        IDENTIFICACIÓN Y PREPARACIÓN DEL PROYECTO:
        - Problema Identificado: {responses.get('problema_identificacion', '')}
        - Análisis de Causas: {responses.get('causas_problema', '')}
        - Análisis de Efectos: {responses.get('efectos_problema', '')}
        - Población Afectada: {responses.get('poblacion_afectada', '')}

        FORMULACIÓN DEL PROYECTO:
        - Objetivo General: {responses.get('objetivo_general', '')}
        - Objetivos Específicos: {responses.get('objetivos_especificos', '')}
        - Justificación: {responses.get('justificacion', '')}
        - Localización: {responses.get('localizacion', '')}
        - Alternativas de Solución: {responses.get('alternativas_solucion', '')}

        COMPONENTES TÉCNICOS:
        - Componentes Técnicos: {responses.get('componentes_tecnicos', '')}
        - Cadena de Valor: {responses.get('cadena_valor', '')}
        - Indicadores: {responses.get('indicadores', '')}

        EVALUACIÓN Y SOSTENIBILIDAD:
        - Análisis de Riesgos: {responses.get('riesgos', '')}
        - Estrategia de Sostenibilidad: {responses.get('sostenibilidad', '')}

        ASPECTOS FINANCIEROS Y CRONOGRAMA:
        - Presupuesto General: {responses.get('presupuesto_general', '')}
        - Presupuesto Detallado: {responses.get('presupuesto_detallado', '')}
        - Cronograma: {responses.get('cronograma', '')}

        Por favor, estructura el documento con las siguientes secciones principales:
        1. Resumen Ejecutivo
        2. Identificación del Problema (Árbol de Problemas)
        3. Objetivos del Proyecto (Árbol de Objetivos)
        4. Justificación y Marco de Referencia
        5. Descripción del Proyecto y Alternativas
        6. Componentes Técnicos IDEC/IA
        7. Cadena de Valor e Indicadores
        8. Análisis de Riesgos y Sostenibilidad
        9. Presupuesto y Cronograma
        10. Conclusiones y Recomendaciones
        """
        
        # Agregar el mensaje al thread
        message = client.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=message_content
        )
        
        # Ejecutar el asistente
        run = client.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=ASSISTANT_ID
        )
        
        # Esperar a que el asistente termine de procesar
        max_attempts = 60  # Máximo 60 segundos
        attempts = 0
        
        while attempts < max_attempts:
            run_status = client.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id
            )
            
            logger.info(f"Estado del asistente: {run_status.status}")
            
            if run_status.status == 'completed':
                break
            elif run_status.status == 'failed':
                raise Exception(f"El asistente falló: {run_status.last_error}")
            elif run_status.status == 'expired':
                raise Exception("El asistente expiró")
            elif run_status.status == 'cancelled':
                raise Exception("El asistente fue cancelado")
            
            time.sleep(1)  # Esperar 1 segundo antes de verificar nuevamente
            attempts += 1
        
        if attempts >= max_attempts:
            raise Exception("Timeout: El asistente tardó demasiado en responder")
        
        # Obtener la respuesta del asistente
        messages = client.beta.threads.messages.list(thread_id=thread.id)
        
        # La respuesta más reciente será la del asistente
        assistant_message = None
        for msg in messages.data:
            if msg.role == "assistant":
                assistant_message = msg
                break
        
        if not assistant_message:
            raise Exception("No se recibió respuesta del asistente")
        
        # Extraer el contenido del mensaje
        content = ""
        if assistant_message.content and len(assistant_message.content) > 0:
            content = assistant_message.content[0].text.value
        
        if not content:
            raise Exception("El asistente devolvió un contenido vacío")
        
        logger.info("Documento generado exitosamente por el asistente")
        return content

    except Exception as e:
        logger.error(f"Error generando documento con asistente: {str(e)}")
        # Intentar con fallback si el asistente falla
        logger.info("Intentando usar fallback con chat completions")
        return generate_document_with_fallback(responses)

def generate_document_with_fallback(responses):
    """Función de fallback usando chat completions directas"""
    try:
        # Crear el prompt para ChatGPT
        prompt = f"""
        Por favor, genera un documento de proyecto de inversión basado en la siguiente información:

        Entidad: {responses.get('entidad_nombre', '')}
        Sector: {responses.get('entidad_sector', '')}
        Componentes IDEC: {responses.get('componentes_idec', '')}
        Problemática: {responses.get('problema_descripcion', '')}
        Situación Actual: {responses.get('situacion_actual', '')}
        Causas y Efectos: {responses.get('causas_efectos', '')}
        Población Afectada: {responses.get('poblacion_afectada', '')}
        Objetivo General: {responses.get('objetivo_general', '')}
        Objetivos Específicos: {responses.get('objetivos_especificos', '')}
        Localización: {responses.get('localizacion', '')}
        Cadena de Valor: {responses.get('cadena_valor', '')}
        Riesgos: {responses.get('riesgos', '')}
        Sostenibilidad: {responses.get('sostenibilidad', '')}
        Presupuesto General: {responses.get('presupuesto_general', '')}
        Presupuesto Detalle: {responses.get('presupuesto_detalle', '')}

        Por favor, genera un documento profesional y bien estructurado que incluya:
        1. Resumen ejecutivo
        2. Información de la entidad
        3. Descripción del problema y situación actual
        4. Objetivos y población objetivo
        5. Localización y cadena de valor
        6. Análisis de riesgos
        7. Estrategia de sostenibilidad
        8. Presupuesto y detalles financieros
        """

        # Llamar a la API de Azure OpenAI usando chat completions
        response = client.chat.completions.create(
            model=os.getenv('OPENAI_DEPLOYMENT_NAME', 'gpt-4o-mini'),
            messages=[
                {"role": "system", "content": "Eres un experto en redacción de proyectos de inversión."},
                {"role": "user", "content": prompt}
            ]
        )

        return response.choices[0].message.content

    except Exception as e:
        logger.error(f"Error generando documento con fallback: {str(e)}")
        raise

def save_document(content):
    try:
        # Crear un nuevo documento
        doc = Document()
        
        # Agregar título
        title = doc.add_heading('Propuesta de Proyecto de Inversión', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar contenido
        doc.add_paragraph(content)
        
        # Guardar el documento
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'proyecto_inversion_{timestamp}.docx'
        filepath = os.path.join(DOCUMENTS_DIR, filename)
        doc.save(filepath)
        
        return filename

    except Exception as e:
        logger.error(f"Error guardando documento: {str(e)}")
        raise

@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        return send_from_directory(DOCUMENTS_DIR, filename, as_attachment=True)
    except Exception as e:
        logger.error(f"Error descargando archivo: {str(e)}")
        return "Error al descargar el archivo", 404

@app.route('/config.json')
def serve_config():
    try:
        config_path = os.path.join(app.static_folder, 'config.json')
        return send_from_directory(app.static_folder, 'config.json')
    except Exception as e:
        logger.error(f"Error sirviendo config.json: {str(e)}")
        return jsonify({'error': 'Configuración no disponible'}), 404

# Generar archivo de configuración para GitHub Pages
def generate_github_pages_config():
    config = {
        'conversation_flow': conversation_flow,
        'api_endpoint': '/api/chat',
        'version': '2.0',
        'description': 'Asistente para formulación de proyectos de inversión MGA/IDEC/IA'
    }
    
    # Crear el archivo en la carpeta static del directorio actual
    config_path = os.path.join(app.static_folder, 'config.json')
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)

if __name__ == '__main__':
    generate_github_pages_config()
    app.run(host='0.0.0.0', port=5001, debug=True) 