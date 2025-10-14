from flask import Flask, render_template, request, jsonify, session, send_from_directory, url_for
from flask_cors import CORS
import os, traceback, logging, json, time, httpx, re
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from openai import AzureOpenAI

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
load_dotenv()

app = Flask(__name__, static_folder='static', template_folder='templates')
CORS(app)
app.secret_key = os.getenv('SECRET_KEY', 'idec_secret_key_change_in_production')

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCUMENTS_DIR = os.path.join(app.static_folder, 'documents')
FORMULARIOS_DIR = os.path.join(app.static_folder, 'formularios')
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(FORMULARIOS_DIR, exist_ok=True)

client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-05-01-preview"),
    http_client=httpx.Client(verify=False)
)

# -------------------------- LLM helper --------------------------
def ask_markdown_azure(messages, *, model_name=None, max_tokens=1500, temperature=0.4, max_rounds=3):
    full_text, rounds = "", 0
    _messages = list(messages)
    model_name = model_name or os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
    while rounds < max_rounds:
        rounds += 1
        resp = client.chat.completions.create(
            model=model_name, messages=_messages, temperature=temperature, max_tokens=max_tokens
        )
        choice = resp.choices[0]
        chunk = (choice.message.content or "").strip()
        full_text += chunk
        finish = getattr(choice, "finish_reason", None)
        if finish not in ("length", "content_filter"):
            break
        _messages += [
            {"role": "assistant", "content": chunk},
            {"role": "user", "content": "Por favor continúa exactamente donde te quedaste."}
        ]
    return full_text

# -------------------------- DOCX helpers --------------------------
def _add_rich_text(paragraph, text: str):
    token_re = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')
    parts = token_re.split(text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1]); run.font.name = "Courier New"; run.font.size = Pt(10)
        else:
            paragraph.add_run(part)

def _add_markdown_line(doc, line: str):
    s = line.strip()
    if not s: return
    if s == '---':
        p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.LINE); return
    if s.startswith('### '): doc.add_heading(s[4:], level=3); return
    if s.startswith('## '):  doc.add_heading(s[3:], level=2); return
    if s.startswith('# '):   doc.add_heading(s[2:], level=1); return
    if re.match(r'^\d+\.\s', s):
        p = doc.add_paragraph(style='List Number')
        _add_rich_text(p, re.sub(r'^\d+\.\s', '', s, 1)); return
    if s.startswith('- ') or s.startswith('* '):
        p = doc.add_paragraph(style='List Bullet')
        _add_rich_text(p, s[2:]); return
    p = doc.add_paragraph(); _add_rich_text(p, s)

def _filtered_responses_for_report(responses: dict) -> dict:
    # No incluir referencias a files subidos
    return {k: v for k, v in responses.items() if not k.startswith('upload_')}

def generate_project_document(responses: dict, filename: str = None) -> str:
    if not filename:
        filename = f"proyecto_inversion_{int(time.time())}.docx"
    filepath = os.path.join(DOCUMENTS_DIR, filename)
    clean = _filtered_responses_for_report(responses)

    prompt = (
        "Eres un experto en formulación de proyectos bajo la MGA (DNP). Organiza la "
        "siguiente información en un proyecto de inversión IDEC/IA. Redacta en español, "
        "Markdown (#, ##, ###, listas). No HTML.\n\n"
        f"JSON: {json.dumps(clean, indent=2, ensure_ascii=False)}\n\n"
        "Incluye secciones: Introducción; Problema/Oportunidad; Población afectada y objetivo; "
        "Localización; Objetivo central; Medios y fines; Cadena de valor; Componentes (si IDEC); Conclusión."
    )
    completion = client.chat.completions.create(
        model=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
        messages=[
            {"role": "system", "content": "Responde en Markdown puro."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2500, temperature=0.5
    )
    md_text = (completion.choices[0].message.content or "").strip()
    doc = Document(); doc.add_heading("Proyecto de Inversión en IDEC/IA", level=0)
    for line in md_text.splitlines(): _add_markdown_line(doc, line)
    doc.save(filepath); return filepath

# -------------------------- Rutas básicas + plantillas + upload --------------------------
@app.route('/')
def index():
    session.clear()
    session['current_step'] = 'intro_bienvenida'
    session['responses'] = {}
    session['mode'] = 'flow'
    return render_template('index.html')

@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        return send_from_directory(DOCUMENTS_DIR, filename, as_attachment=True)
    except Exception as e:
        logger.error(f"Error descargando archivo: {e}")
        return "Error al descargar el archivo", 404

@app.route('/plantilla/<tipo>')
def plantilla(tipo):
    fname = 'PlantillaCausa.xlsx' if tipo == 'causa' else 'PlantillaObjetivo.xlsx' if tipo == 'objetivo' else None
    if not fname: return "Tipo inválido", 400
    path = os.path.join(BASE_DIR, fname)
    if not os.path.isfile(path): return "Plantilla no encontrada", 404
    return send_from_directory(BASE_DIR, fname, as_attachment=True)

@app.route('/api/upload_formulario', methods=['POST'])
def upload_formulario():
    if 'file' not in request.files:
        return jsonify({"ok": False, "error": "Archivo no recibido"}), 400
    f = request.files['file']
    tipo = (request.form.get('tipo') or '').strip().lower()
    if tipo not in ('causa', 'objetivo'):
        return jsonify({"ok": False, "error": "Tipo inválido"}), 400

    proyecto = session.get('responses', {}).get('nombre_proyecto', 'proyecto')
    slug = re.sub(r'[^A-Za-z0-9_\-]+', '_', proyecto).strip('_') or 'proyecto'
    filename = f"{tipo}-{slug}.xsls"  # extensión solicitada

    save_path = os.path.join(FORMULARIOS_DIR, filename)
    f.save(save_path)

    resp = session.get('responses', {})
    resp[f"upload_{tipo}"] = filename
    session['responses'] = resp

    return jsonify({"ok": True, "filename": filename})

@app.route('/reset', methods=['POST'])
def reset_conversation():
    session.clear()
    session['current_step'] = 'intro_bienvenida'
    session['responses'] = {}
    session['mode'] = 'flow'
    return jsonify({"status": "ok", "message": "Conversación reiniciada"})

# -------------------------- Chat Libre --------------------------
def _md_link(url, text): return f"[{text}]({url})"
def _is_yes(txt): return bool(re.search(r'\b(sí|si)\b', txt))
def _is_no(txt):  return bool(re.search(r'\bno\b', txt))

def _bootstrap_alt_explanation(topic_md: str):
    session['mode'] = 'alt'
    system_msg = {"role": "system", "content": "Responde SIEMPRE en Markdown claro, con viñetas y ejemplo."}
    user_msg   = {"role": "user", "content": f"{topic_md}\n\nTermina con: 'Cuando estés listo, escribe **Finalizar** para volver al flujo.'"}
    md = ask_markdown_azure([system_msg, user_msg], max_tokens=1000, temperature=0.4)
    return "💬 Has activado el **Chat Libre** para resolver esta duda.\n\n" + md

@app.route('/api/chat_alt', methods=['POST'])
def chat_alt():
    try:
        data = request.get_json(silent=True) or {}
        user_message = (data.get("message") or "").strip()
        session['mode'] = "alt"

        if user_message.lower() == "finalizar":
            session['mode'] = "flow"
            next_after = session.pop('after_alt_next_step', None)
            if next_after: session['current_step'] = next_after
            session['resume_from_alt'] = True
            return jsonify({
                "response": "✅ Has finalizado el chat libre. Volvemos al flujo normal.",
                "options": ["Continuar flujo"],
                "format": "markdown"
            })

        md = ask_markdown_azure(
            [{"role":"system","content":"Responde en Markdown válido, sin HTML."},
             {"role":"user","content":user_message}],
            model_name=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
            max_tokens=1500, temperature=0.4, max_rounds=3
        )
        return jsonify({"response": md, "format": "markdown"})
    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500

# -------------------------- Flujo (preguntas) --------------------------
conversation_flow = {
    "intro_bienvenida": {
        "prompt":
            "👋 ¡Hola! Soy tu asistente virtual para ayudarte en la formulación de proyectos de inversión relacionados con Infraestructura de Datos (IDEC) o Inteligencia Artificial (IA). Vamos a empezar paso a paso.\n\n"
            "Te acompañaré paso a paso para estructurar tu proyecto conforme a la Metodología General Ajustada (MGA) del Departamento Nacional de Planeación.\n\n"
            "🧰 Te haré preguntas clave para estructurar el proyecto.\n\n"
            "❓ Antes de continuar, ¿todo está claro? o ¿tienes algunas preguntas?",
        "options": [
            "Sí, entiendo el proceso y deseo continuar",
            "Tengo dudas respecto al proceso, me gustaría resolverlas antes de empezar"
        ],
        "next_step": "pregunta_3_entidad"
    },
    "gate_1_ciclo": {
        "prompt": "🔎 ¿Conoces el ciclo de inversión pública y las fases que lo componen?",
        "options": ["Sí, lo conozco", "No, no lo conozco"],
        "next_step": "gate_2_herramienta"
    },
    "gate_2_herramienta": {
        "prompt": "🧭 ¿Comprende que esta herramienta es de orientación y que el borrador resultante puede emplearse como insumo o apoyo en la etapa de formulación?",
        "options": ["Sí, lo comprendo", "No, no lo tengo claro"],
        "next_step": "pregunta_3_entidad"
    },

    "pregunta_3_entidad": {"prompt": "🏢 ¿Cuál es el nombre de tu entidad?", "next_step": "rol_abierto"},
    "rol_abierto": {
        "prompt": "👤 ¿Cuál es su rol dentro de la entidad (por ejemplo: Director de área, Coordinador, Profesional especializado, Analista, Asesor, Técnico operativo, Contratista de apoyo)?",
        "next_step": "elige_vertical"
    },

    "elige_vertical": {
        "prompt": "💡 ¿Deseas construir un proyecto de inversión asociando componentes de tecnologías de la información y las comunicaciones en temas de Infraestructura de datos (IDEC) o Inteligencia Artificial (IA)?",
        "options": ["Sí, en IDEC", "Sí, en IA", "No (Cierre de la conversación)"],
        "next_step": "nombre_proyecto"
    },

    "idec_componentes": {
        "prompt":
            "📚 La siguiente es la lista de los componentes que integran la IDEC, por favor selecciona los componentes que deseas incluir en tu proyecto de inversión. Selección múltiple :\n",
        "next_step": "nombre_proyecto"
    },

    "nombre_proyecto": {"prompt": "📝 ¿Cuál es el nombre del proyecto de inversión?", "next_step": "poblacion_afectada"},
    "poblacion_afectada": {"prompt": "👥 ¿Cuál es la población afectada por el proyecto de inversión? Descríbela y asocia un número", "next_step": "poblacion_objetivo"},
    "poblacion_objetivo": {"prompt": "🎯 ¿Cuál es la población objetivo que pretende ser beneficiada de la intervención que realiza el proyecto de inversión? Descríbela y asocia un número", "next_step": "localizacion"},
    "localizacion": {"prompt": "📍 ¿Cuál es la localización en la que se enmarca el proyecto (Ejemplo: Territorial-Territorio Norte, nacional-Colombia, departamental-Cundinamarca)?", "next_step": "problema_oportunidad"},
    "problema_oportunidad": {"prompt": "🧩 ¿Cuál es la problemática o la oportunidad que tu proyecto de inversión busca atender o resolver?", "next_step": "upload_causas"},

    "upload_causas": {"prompt": "📄 Cargue la plantilla diligenciada con las causas estructuradas. Recuerde que cada causa debe incluir dos causas indirectas, un efecto directo y un efecto indirecto.", "next_step": "upload_objetivos"},
    "upload_objetivos": {"prompt": "🎯 Cargue la plantilla diligenciada con los objetivos estructurados. Recuerde que cada objetivo debe incluir un medio directo, al menos un medio indirecto, un fin directo y un fin indirecto.", "next_step": "cadena_valor"},

    "cadena_valor": {"prompt": "🔗 ¿Cómo se constituye tu cadena de valor?", "next_step": "finalizado"}
}

def _upload_prompt_with_link(step_key: str) -> str:
    if step_key == 'upload_causas':
        url = url_for('plantilla', tipo='causa')
        return f"**Cargar plantilla de causas.**\n\n1. Descargue la plantilla.\n2. Diligénciela.\n3. Súbala en el recuadro que aparece debajo.\n\n**Descargar plantilla:** [PlantillaCausa.xlsx]({url})"
    else:
        url = url_for('plantilla', tipo='objetivo')
        return f"**Cargar plantilla de objetivos.**\n\n1. Descargue la plantilla.\n2. Diligénciela.\n3. Súbala en el recuadro que aparece debajo.\n\n**Descargar plantilla:** [PlantillaObjetivo.xlsx]({url})"

# -------------------------- Handler principal --------------------------
@app.route('/api/chat', methods=['POST'])
def chat():
    # Si estamos en chat libre, delegar
    if session.get("mode") == "alt":
        return chat_alt()

    data = request.get_json() or {}
    user_message = (data.get('message') or '').strip()
    user_lower = user_message.lower()

    current_step = session.get('current_step', 'intro_bienvenida')
    responses = session.get('responses', {})

    # Inicio
    if current_step == 'intro_bienvenida' and user_lower in ('iniciar', 'start'):
        intro = conversation_flow['intro_bienvenida']
        return jsonify({"response": intro['prompt'], "current_step": "intro_bienvenida", "options": intro.get('options', []), "format": "markdown"})

    # Opción dudas
    if current_step == 'intro_bienvenida' and user_lower == 'tengo dudas respecto al proceso, me gustaría resolverlas antes de empezar':
        session['current_step'] = 'gate_1_ciclo'
        step = conversation_flow['gate_1_ciclo']
        return jsonify({"response": step['prompt'], "current_step": "gate_1_ciclo", "options": step['options'], "format": "markdown"})

    if current_step == 'intro_bienvenida' and user_lower == 'sí, entiendo el proceso y deseo continuar':
        session['current_step'] = conversation_flow['intro_bienvenida']['next_step']
        step = conversation_flow[session['current_step']]
        payload = {"response": step['prompt'], "current_step": session['current_step'], "format": "markdown"}
        if "options" in step: payload["options"] = step["options"]
        return jsonify(payload)

    # Reanudar VOLVER DEL CHAT LIBRE
    if session.pop('resume_from_alt', False) or user_lower in ('continuar flujo', 'volver al flujo'):
        step_key = session.get('current_step', 'intro_bienvenida')
        step_conf = conversation_flow.get(step_key, {})
        resp_text = step_conf.get("prompt", "…")

        # Reinyectar tarjetas si corresponde
        if step_key == 'idec_componentes':
            step = conversation_flow['idec_componentes']
            return jsonify({
                "response": step['prompt'] + "\n\nSelecciona una o más tarjetas y pulsa **Confirmar**.",
                "current_step": "idec_componentes",
                "format": "markdown",
                "multiselect": {
                    "items": [
                        "Gobernanza de datos",
                        "Interoperabilidad",
                        "Herramientas técnicas y tecnológicas",
                        "Seguridad y privacidad de datos",
                        "Datos",
                        "Aprovechamiento de datos"
                    ],
                    "submit_text": "Confirmar"
                }
            })

        if step_key in ('upload_causas','upload_objetivos'):
            resp_text = _upload_prompt_with_link(step_key)
            tipo = 'causa' if step_key == 'upload_causas' else 'objetivo'
            return jsonify({
                "response": resp_text, "current_step": step_key, "format": "markdown",
                "upload": {"expect_upload": True, "tipo": tipo, "download_url": url_for('plantilla', tipo=tipo)}
            })
        payload = {"response": resp_text, "current_step": step_key, "format": "markdown"}
        if "options" in step_conf: payload["options"] = step_conf["options"]
        return jsonify(payload)

    # Gates
    if current_step == 'gate_1_ciclo':
        if _is_yes(user_lower):
            session['current_step'] = conversation_flow['gate_1_ciclo']['next_step']
            step = conversation_flow['gate_2_herramienta']
            return jsonify({"response": step['prompt'], "current_step": "gate_2_herramienta", "options": step['options'], "format": "markdown"})
        elif _is_no(user_lower):
            session['after_alt_next_step'] = "gate_2_herramienta"
            md = _bootstrap_alt_explanation("Explica el ciclo de inversión pública y sus fases principales.")
            return jsonify({"response": md, "format": "markdown"})
        else:
            step = conversation_flow['gate_1_ciclo']
            return jsonify({"response": step['prompt'], "current_step": "gate_1_ciclo", "options": step['options'], "format": "markdown"})

    if current_step == 'gate_2_herramienta':
        if _is_yes(user_lower):
            session['current_step'] = conversation_flow['gate_2_herramienta']['next_step']
            step = conversation_flow[session['current_step']]
            payload = {"response": step['prompt'], "current_step": session['current_step'], "format": "markdown"}
            if "options" in step: payload["options"] = step["options"]
            return jsonify(payload)
        elif _is_no(user_lower):
            session['after_alt_next_step'] = "pregunta_3_entidad"
            md = _bootstrap_alt_explanation("Explica por qué esta herramienta es de orientación y cómo el borrador sirve como insumo en formulación (MGA).")
            return jsonify({"response": md, "format": "markdown"})
        else:
            step = conversation_flow['gate_2_herramienta']
            return jsonify({"response": step['prompt'], "current_step": "gate_2_herramienta", "options": step['options'], "format": "markdown"})

    # Registro inicial
    if current_step == 'rol_abierto' and user_message:
        responses[current_step] = user_message
        session['responses'] = responses
        session['current_step'] = 'elige_vertical'
        step = conversation_flow['elige_vertical']
        return jsonify({"response": step['prompt'], "current_step": "elige_vertical", "options": step['options'], "format": "markdown"})

    # Elegir vertical
    if current_step == 'elige_vertical':
        choice = user_lower.strip()
        if 'cerrar la conversación' in choice or choice == 'no' or choice.startswith('no '):
            session['current_step'] = 'finalizado'
            msg = "❌ Este asistente solo atiende proyectos **IDEC/IA**. Se cierra la conversación. Usa *Reiniciar* para empezar de nuevo."
            return jsonify({"response": msg, "current_step": "finalizado", "format": "markdown"})
        elif 'idec' in choice:
            responses['vertical'] = 'IDEC'; session['responses'] = responses
            session['current_step'] = 'idec_componentes'
            step = conversation_flow['idec_componentes']
            return jsonify({
                "response": step['prompt'] + "\n\nSelecciona una o más tarjetas y pulsa **Confirmar**.",
                "current_step": "idec_componentes",
                "format": "markdown",
                "multiselect": {
                    "items": [
                        "Gobernanza de datos",
                        "Interoperabilidad",
                        "Herramientas técnicas y tecnológicas",
                        "Seguridad y privacidad de datos",
                        "Datos",
                        "Aprovechamiento de datos"
                    ],
                    "submit_text": "Confirmar"
                }
            })
        elif 'ia' in choice:
            responses['vertical'] = 'IA'; session['responses'] = responses
            session['current_step'] = conversation_flow['elige_vertical']['next_step']  # nombre_proyecto
            step = conversation_flow[session['current_step']]
            return jsonify({"response": step['prompt'], "current_step": session['current_step'], "format": "markdown"})
        else:
            step = conversation_flow['elige_vertical']
            return jsonify({"response": step['prompt'], "current_step": "elige_vertical", "options": step['options'], "format": "markdown"})

    # IDEC multiselección (tarjetas largas apiladas)
    if current_step == 'idec_componentes':
        if user_message.startswith('__msel__:'):
            raw = user_message.split(':', 1)[1]
            comps = [c.strip() for c in raw.split('|') if c.strip()]
            if comps:
                responses['idec_componentes'] = comps
                session['responses'] = responses
                session['current_step'] = conversation_flow['idec_componentes']['next_step']  # nombre_proyecto
                step = conversation_flow[session['current_step']]
                return jsonify({"response": step['prompt'], "current_step": session['current_step'], "format": "markdown"})
        # Re-mostrar selecciones si no vino en el formato esperado
        step = conversation_flow['idec_componentes']
        return jsonify({
            "response": step['prompt'] + "\n\nSelecciona una o más tarjetas y pulsa **Confirmar**.",
            "current_step": "idec_componentes",
            "format": "markdown",
            "multiselect": {
                "items": [
                    "Gobernanza de datos",
                    "Interoperabilidad",
                    "Herramientas técnicas y tecnológicas",
                    "Seguridad y privacidad de datos",
                    "Datos",
                    "Aprovechamiento de datos"
                ],
                "submit_text": "Confirmar"
            }
        })

    # ----- PASOS DE CARGA: permitir subir y avanzar con "Continuar" -----
    if current_step in ('upload_causas', 'upload_objetivos'):
        step_key = current_step
        tipo = 'causa' if step_key == 'upload_causas' else 'objetivo'
        required_flag = f"upload_{tipo}"

        if re.search(r'\b(continuar|siguiente)\b', user_lower):
            if required_flag in session.get('responses', {}):
                next_step = conversation_flow[step_key]['next_step']
                session['current_step'] = next_step
                step_conf = conversation_flow[next_step]
                text = step_conf.get("prompt", "…")
                payload = {"response": text, "current_step": next_step, "format": "markdown"}
                if "options" in step_conf: payload["options"] = step_conf["options"]
                if next_step in ('upload_causas', 'upload_objetivos'):
                    next_tipo = 'causa' if next_step == 'upload_causas' else 'objetivo'
                    payload["response"] = _upload_prompt_with_link(next_step)
                    payload["upload"] = {"expect_upload": True, "tipo": next_tipo, "download_url": url_for('plantilla', tipo=next_tipo)}
                return jsonify(payload)
            else:
                text = _upload_prompt_with_link(step_key) + "\n\n> ⚠️ Aún no has subido el archivo. Por favor súbelo y luego escribe **Continuar**."
                return jsonify({
                    "response": text, "current_step": step_key, "format": "markdown",
                    "upload": {"expect_upload": True, "tipo": tipo, "download_url": url_for('plantilla', tipo=tipo)}
                })

        text = _upload_prompt_with_link(step_key)
        return jsonify({
            "response": text, "current_step": step_key, "format": "markdown",
            "upload": {"expect_upload": True, "tipo": tipo, "download_url": url_for('plantilla', tipo=tipo)}
        })

    # Guardar y avanzar (genérico)
    responses[current_step] = user_message
    session['responses'] = responses

    next_step = conversation_flow.get(current_step, {}).get("next_step")
    if (not next_step) or (next_step == "finalizado"):
        session['current_step'] = "finalizado"
        filepath = generate_project_document(responses)
        filename = os.path.basename(filepath)
        md_link = _md_link(url_for('download_file', filename=filename), "Descargar documento")
        return jsonify({"response": f"✅ Flujo completado. Documento generado. {md_link}", "current_step": "finalizado", "format": "markdown"})

    session['current_step'] = next_step
    step_conf = conversation_flow.get(next_step, {})
    text = step_conf.get("prompt", "…")
    payload = {"response": text, "current_step": next_step, "format": "markdown"}
    if "options" in step_conf: payload["options"] = step_conf["options"]
    if next_step in ('upload_causas','upload_objetivos'):
        next_tipo = 'causa' if next_step == 'upload_causas' else 'objetivo'
        payload["response"] = _upload_prompt_with_link(next_step)
        payload["upload"]  = {"expect_upload": True, "tipo": next_tipo, "download_url": url_for('plantilla', tipo=next_tipo)}
    return jsonify(payload)

# -------------------------- Run --------------------------
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True)
