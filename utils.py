
# utils.py
# ============================================================
# Utilidades para el chatbot IDEC/IA:
# - LLM helper (Azure OpenAI)
# - Generación de DOCX con secciones ordenadas y títulos (sin mostrar IDs)
# - Validadores + Parsers de plantillas Excel
# - Guardado y carga de árboles JSON (UTF-8 con BOM)
# - Conversation flow (para importar desde app.py)
# ============================================================

from __future__ import annotations
import os
import re
import json
import time
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from openpyxl import load_workbook
from datetime import datetime
import pandas as pd

# Fecha actual
fecha = datetime.now()
meses = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}
fecha_actual = f"{fecha.day} de {meses[fecha.month]} de {fecha.year}"




SYSTEM_PRIMER = """
Contexto fijo:
- País por defecto: Colombia. Cuando se hable de departamentos/municipios/localidades, se asume Colombia.
- DNP = Departamento Nacional de Planeación (Colombia).
- IDEC = Infraestructura de Datos del Estado Colombiano.
- Usa terminología y normatividad de Colombia cuando aplique.
- Si te dan porcentajes o proporciones sin base absoluta, explica el cálculo y estima usando datos oficiales si están disponibles.
- No muestres códigos internos de árbol (C1, CI1, O1, MI1) en el texto final.
"""

# -------------------------- LLM helper --------------------------
def ask_markdown_azure(
    messages: List[Dict[str, str]],
    *,
    client,
    model_name: Optional[str] = None,
    max_tokens: int = 1800,
    temperature: float = 0.4,
    max_rounds: int = 3,
    use_primer = True
) -> str:
    """Envía mensajes a Azure OpenAI y concatena si se corta por longitud."""
    full_text, rounds = "", 0
    _messages = list(messages)
    if use_primer:
        sys = {"role": "system", "content": SYSTEM_PRIMER + "\nResponde en Markdown válido."}
        _messages = [sys] + _messages
    if model_name is None:
        model_name = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
    while rounds < max_rounds:
        rounds += 1
        resp = client.chat.completions.create(
            model=model_name, messages=_messages, temperature=temperature, max_tokens=max_tokens
        )
        choice = resp.choices[0]
        chunk = (choice.message.content or "").strip()
        full_text += chunk
        finish = getattr(choice, "finish_reason", None)
        if finish not in ("length", "content_filter"):
            break
        _messages += [
            {"role": "assistant", "content": chunk},
            {"role": "user", "content": "Por favor continúa exactamente donde te quedaste."},
        ]
    return full_text


# -------------------------- DOCX helpers --------------------------
def _add_rich_text(paragraph, text: str) -> None:
    """Aplica **negrita**, *itálica*, `monoespaciado` y [enlaces](url) simple dentro de un párrafo."""
    # Procesar enlaces primero: [texto](url) -> texto (url)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    
    # Procesar texto con negrita, itálica y monoespaciado
    token_re = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(part)


def _add_markdown_line(doc, line: str) -> None:
    """Convierte una línea de Markdown muy simple a estructuras de docx.
    Soporta #, ##, ###, ####; listas numeradas y con viñetas.
    """
    s = line.strip()
    if not s:
        return
    if s == '---':
        p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.LINE); return
    if s.startswith('#### '):
        doc.add_heading(s[5:], level=4); return
    if s.startswith('### '):
        doc.add_heading(s[4:], level=3); return
    if s.startswith('## '):
        doc.add_heading(s[3:], level=2); return
    if s.startswith('# '):
        doc.add_heading(s[2:], level=1); return
    if re.match(r'^\d+\.\s', s):
        p = doc.add_paragraph(style='List Number'); _add_rich_text(p, re.sub(r'^\d+\.\s', '', s, 1)); return
    if s.startswith('- ') or s.startswith('* '):
        p = doc.add_paragraph(style='List Bullet'); _add_rich_text(p, s[2:]); return
    p = doc.add_paragraph(); _add_rich_text(p, s)


def _filtered_responses_for_report(responses: dict) -> dict:
    """Filtra claves internas (e.g., uploads) para el reporte."""
    return {k: v for k, v in responses.items() if not k.startswith('upload_')}


# -------------------------- Árbol -> Outline para prompt --------------------------
def causas_tree_to_outline(tree: Dict[str, Any]) -> str:
    """Devuelve un outline sin códigos (C1, CI1, etc.)."""
    if not tree or not tree.get("items"): 
        return "(sin causas)"
    lines = ["Marco del problema: Causas y efectos"]
    for c in tree["items"]:
        cdesc = (c.get("descripcion") or "").strip()
        edesc = ((c.get("efecto_directo") or {}).get("descripcion") or "").strip()
        lines.append(f"Causa: {cdesc}")
        if edesc:
            lines.append(f"Efecto directo: {edesc}")
        cis = c.get("causas_indirectas", [])
        if cis:
            lines.append("Causas indirectas:")
            for ci in cis:
                cidesc = (ci.get("descripcion") or "").strip()
                lines.append(f"  a) {cidesc}")
                for ei in ci.get("efectos_indirectos", []):
                    lines.append(f"     * Efecto indirecto: {(ei.get('descripcion') or '').strip()}")
    return "\n".join(lines)


def objetivos_tree_to_outline(tree: Dict[str, Any]) -> str:
    """Devuelve un outline sin códigos (O1, MI1, etc.)."""
    if not tree or not tree.get("items"): 
        return "(sin objetivos)"
    lines = ["Marco de objetivos: Medios y fines"]
    for o in tree["items"]:
        odesc = (o.get("descripcion") or "").strip()
        md = ((o.get("medio_directo") or {}).get("descripcion") or "").strip()
        fd = ((o.get("fin_directo") or {}).get("descripcion") or "").strip()
        lines.append(f"Objetivo: {odesc}")
        if md: lines.append(f"Medio directo: {md}")
        if fd: lines.append(f"Fin directo: {fd}")
        mis = o.get("medios_indirectos", [])
        if mis:
            lines.append("Medios indirectos y fines:")
            for mi in mis:
                midesc = (mi.get("descripcion") or "").strip()
                lines.append(f"  a) {midesc}")
                for fi in mi.get("fines_indirectos", []):
                    lines.append(f"     * Fin indirecto: {(fi.get('descripcion') or '').strip()}")
    return "\n".join(lines)


# -------------------------- Carga/guardado JSON --------------------------
def load_tree_json(path: str) -> Optional[Dict[str, Any]]:
    try:
        with open(path, "r", encoding="utf-8-sig") as f:
            return json.load(f)
    except Exception:
        return None


def save_tree_json(tree: Dict[str, Any], out_dir: str, base_filename: str, *, encoding: str = "utf-8-sig") -> str:
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{base_filename}.json")
    with open(out_path, "w", encoding=encoding) as f:
        json.dump(tree, f, ensure_ascii=False, indent=2)
    return out_path


# -------------------------- Generación de documento --------------------------
def generate_project_document(
    responses: dict,
    *,
    client,
    documents_dir: str,
    filename: Optional[str] = None,
    causas_tree: Optional[Dict[str, Any]] = None,
    objetivos_tree: Optional[Dict[str, Any]] = None,
    formularios_json_dir: Optional[str] = None,
) -> str:
    """Genera el .docx del proyecto con secciones que justifican el proyecto basado en
    los árboles de Causas/Efectos y Objetivos/Medios/Fines, manteniendo el orden de secciones definido.
    """
    if not filename:
        filename = f"proyecto_inversion_{int(time.time())}.docx"
    os.makedirs(documents_dir, exist_ok=True)
    filepath = os.path.join(documents_dir, filename)

    # Cargar árboles desde disco si no vienen en memoria
    if formularios_json_dir:
        # Para plantilla general, buscar archivo JSON único que contiene todo
        if causas_tree is None or objetivos_tree is None:
            if responses.get("upload_plantilla"):
                # El archivo JSON tiene el mismo nombre base que el Excel pero con extensión .json
                # Ejemplo: plantilla-mi-proyecto.xlsx -> plantilla-mi-proyecto.json
                base_plantilla = os.path.splitext(responses["upload_plantilla"])[0]  # sin .xlsx
                json_path = os.path.join(formularios_json_dir, f"{base_plantilla}.json")
                if os.path.exists(json_path):
                    # El JSON contiene todas las hojas con causas y objetivos
                    tree_data = load_tree_json(json_path)
                    if tree_data:
                        # Usar el mismo árbol para causas y objetivos (contiene todo)
                        if causas_tree is None:
                            causas_tree = tree_data
                        if objetivos_tree is None:
                            objetivos_tree = tree_data
            elif responses.get("upload_causa"):
                base = os.path.splitext(responses["upload_causa"])[0]  # sin .xlsx
                if causas_tree is None:
                    causas_tree = load_tree_json(os.path.join(formularios_json_dir, f"{base}.json"))
            elif responses.get("upload_objetivo"):
                base = os.path.splitext(responses["upload_objetivo"])[0]
                if objetivos_tree is None:
                    objetivos_tree = load_tree_json(os.path.join(formularios_json_dir, f"{base}.json"))

    clean = _filtered_responses_for_report(responses)
    causas_outline = causas_tree_to_outline(causas_tree) if causas_tree else "(sin causas)"
    objetivos_outline = objetivos_tree_to_outline(objetivos_tree) if objetivos_tree else "(sin objetivos)"

    # Prompt con orden de secciones fijo y sin códigos de IDs visibles
    prompt = (
    "Eres un experto en formulación de proyectos bajo la Metodología General Ajustada (MGA) del Departamento Nacional de Planeación en Colombia (DNP). "
    "Redacta en ESPAÑOL y devuelve el contenido en Markdown estructurado con #, ##, ### y #### (sin códigos C1/O1 visibles ni siglas sin desarrollar). "
    "El sistema convertirá luego a Word con títulos y estilos formales.\n\n"
    
    "AL INICIO DEL DOCUMENTO, GENERA EL SIGUIENTE ENCABEZADO INSTITUCIONALCENTRADO:\n"
    f"**{fecha_actual}**"
    "\n\n"
    
    "ORDEN OBLIGATORIO DE SECCIONES (usa encabezados Markdown):\n"
    "## Introducción\n"
    "## Planteamiento del problema u oportunidad\n"
    "## Localización\n"
    "## Marco del problema: Causas y efectos\n"
    "## Marco de objetivos: Medios y fines\n"
    "## Componentes del proyecto\n"
    "## Cadena de valor\n"
    "## Conclusión y justificación final\n\n"
    
    "INSTRUCCIONES:\n"
    "- Integra los datos del usuario y los árboles provistos.\n"
    "- No uses siglas ni abreviaturas: escribe los nombres completos de las entidades (por ejemplo, 'Ministerio de Educación Nacional' en lugar de 'MinEducación').\n"
    "- En 'Marco del problema: Causas y efectos': para cada causa, usa '### Causa' con una explicación; luego '#### Efecto directo' y '#### Causas indirectas'.\n"
    "- En 'Marco de objetivos: Medios y fines': usa '### Objetivo', '#### Medio directo', '#### Fin directo' y '#### Medios indirectos'.\n"
    "- En 'Componentes del proyecto' enumera los componentes seleccionados por el usuario y explica brevemente su papel.\n"
    "- Mantén coherencia narrativa entre el problema y los objetivos, y finaliza con una conclusión justificativa del proyecto.\n\n"
    "FORMATO:\n"
    "- El título principal y los datos de encabezado deben ir centrados.\n"
    "- Todo el cuerpo del texto debe estar con alineación justificada.\n\n"

    f"Datos del usuario (JSON):\n{json.dumps(clean, ensure_ascii=False, indent=2)}\n\n"
    "Árbol de causas/efectos (outline):\n" + causas_outline + "\n\n"
    "Árbol de objetivos/medios/fines (outline):\n" + objetivos_outline + "\n\n"
    "RECUERDA: No incluyas códigos como C1, CI1, O1, MI1 en los títulos ni en el texto. "
    "Verifica consistencia numérica, define términos confusos y resume los hallazgos clave al final."
    "En caso de que no te den algunos datos, pero lo puedas conseguir en internet colocalos y referencialos. Por ejemplo, la cantidad de habitantes de alguna zona, si te dan especificaciones de dónde está la pobklación y quiénes son, puedes buscar en tu base de datos o en internet para averiguar qué numero puede ser, estimandolo"
)


    completion = client.chat.completions.create(
        model=os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME"),
        messages=[
            {"role": "system", "content": SYSTEM_PRIMER + "\nResponde exclusivamente en Markdown válido."},
            {"role": "user", "content": prompt},
        ],
        max_tokens=3000,
        temperature=0.4,
    )
    md_text = (completion.choices[0].message.content or "").strip()

    # Escribir DOCX desde Markdown simple
    doc = Document()
    # Título del documento (nivel 0)
    titulo = responses.get("nombre_proyecto") or "Proyecto de Inversión - IDEC/IA"
    doc.add_heading(titulo, level=0)
    
    # Texto aclaratorio y recomendaciones que siempre va después del título
    nota_aclara_md = (
        "**Nota aclaratoria:** Esta plantilla es bosquejo preliminar para la estructuración del proyecto de inversión. "
        "Recordar que esta información debe ser validada y trabajada por la entidad pública, dado que no se constituye "
        "como un documento formal para ser presentado ante la Dirección de Inversiones.\n\n\n"
        "## Recomendaciones\n\n"
        "También con el ánimo de fortalecer el documento que se está construyendo se sugiere revisar las guías y documentos "
        "oficiales sobre formulación de proyectos de inversión, en especial:\n\n"
        "El Manual de usuario del asistente que lo encuentras en el botón de \"Manual de usuario\"\n\n\n"
        "Manuales: Metodología General Ajustada para la formulación de proyectos de inversión pública en Colombia; "
        "Guía orientadora para la definición de productos: "
        "[Manuales DNP](https://www.dnp.gov.co/LaEntidad_/subdireccion-general-inversiones-seguimiento-evaluacion/direccion-proyectos-informacion-para-inversion-publica/Paginas/manuales.aspx)\n\n\n"
        "Cadena de valor: Guía de Cadena de Valor\n\n"
        "Guía para la formulación de indicadores: Guía Metodológica para la formulación de indicadores\n\n"
        "Instrumento de la MGA que consiste en la estandarización de los bienes y servicios que se pueden financiar y generar "
        "a través de los recursos públicos que son ejecutados a través de los proyectos de inversión pública. En este archivo "
        "encontrará la información estandarizada a nivel de sectores, programas y subprogramas; sectores; y productos: "
        "[Catálogo de Productos](https://colaboracion.dnp.gov.co/CDT/proyectosinformacioninversionpublica/catalogos/CATALOGO_DE_PRODUCTOS.xlsx?Web=1)\n\n\n"
        "Las guías de recomendaciones para la formulación de proyectos de inversión de la IDEC e IA (Pendiente ruta)\n\n"
        "Guía de recomendaciones para la formulación de proyectos IDEC e IA para las entidades territoriales: (Pendiente ruta)\n\n\n"
    )
    
    # Agregar el texto aclaratorio al documento
    for line in nota_aclara_md.splitlines():
        _add_markdown_line(doc, line)
    
    # Agregar el contenido generado por la IA
    for line in md_text.splitlines():
        _add_markdown_line(doc, line)
    
    # Texto final que siempre va al final del documento
    texto_final_md = (
        "\n\n"
        "Tener en cuenta que las siguientes secciones deben completarse en el documento final de proyectos de inversión, "
        "dado que este documento es solo un bosquejo preliminar para la estructuración del proyecto de inversión.\n\n\n"
        "En la plantilla que se descargue se incorporen elementos adicionales (vacíos) que debe tener el proyecto:\n\n\n"
        "## Participantes\n\n"
        "- Identificación de los participantes\n"
        "- Análisis de los participantes\n\n"
        "## Población\n\n"
        "- Población afectada por el problema\n"
        "- Población objetivo de la intervención\n\n"
        "## Alternativas de la solución\n\n"
        "- Soluciones identificadas\n"
        "- Alternativa de solución seleccionada\n\n"
        "## Estudio de necesidades\n\n"
        "- Bien o servicio a entregar o demanda a satisfacer\n"
        "- Análisis técnico de la alternativa\n"
        "- Localización de la alternativa\n\n"
        "## Localización\n\n"
        "Localización (Región-Departamento-Municipio-Tipo de agrupación-Agrupación-Específica-Latitud-Longitud)\n\n"
        "## Cadena de valor\n\n"
        "Estructura del Enfoque de Marco Lógico en la cadena de valor con el desarrollo metodológico de las actividades:\n\n"
        "- Producto\n"
        "- Entregable\n"
        "- Indicador\n"
        "- Actividad\n\n"
        "## Análisis de riesgos\n\n"
        "Análisis de riesgos para la alternativa de solución seleccionada\n\n"
        "## Análisis de cuantificación\n\n"
        "Análisis de cuantificación de los ingresos y beneficios\n\n"
        "## Análisis de la estrategia de sostenibilidad\n\n"
        "Análisis de la estrategia de sostenibilidad de la alternativa seleccionada\n\n"
        "## Regionalización de recursos\n\n"
        "Regionalización de recursos (si aplica)\n\n"
        "## Focalización de políticas transversales\n\n"
        "Focalización de políticas transversales (si aplica)\n\n"
        "### Resumen políticas con característica poblacional\n\n"
        "- Políticas con población\n"
        "- Políticas sin población\n"
        "- Cruce de políticas\n"
        "- Resumen de focalización\n"
    )
    
    # Agregar el texto final al documento
    for line in texto_final_md.splitlines():
        _add_markdown_line(doc, line)
    
    doc.save(filepath)
    return filepath


# -------------------------- Utilidades varias --------------------------
def _md_link(url: str, text: str) -> str:
    return f"[{text}]({url})"


def _is_yes(txt: str) -> bool:
    return bool(re.search(r"\b(sí|si)\b", txt or "", flags=re.I))


def _is_no(txt: str) -> bool:
    return bool(re.search(r"\bno\b", txt or "", flags=re.I))


def _num_from_id(id_str: str) -> int:
    """Convierte ID tipo 'C1' o 'O3' a número para ordenar de forma estable."""
    if not id_str:
        return 999999
    digits = ''.join(ch for ch in id_str if ch.isdigit())
    return int(digits) if digits else 999999


def split_sheet_blocks(df: pd.DataFrame):
    """
    Divide automáticamente la hoja en dos bloques:
    - CAUSAS: columnas 0–10
    - OBJETIVOS: columnas 11–22
    """
    CAUSAS_COLS = list(range(0, 11))
    OBJ_COLS = list(range(11, 23))

    df_causas = df.iloc[:, CAUSAS_COLS].dropna(how="all")
    df_obj = df.iloc[:, OBJ_COLS].dropna(how="all")

    return df_causas, df_obj



# -------------------------- Parsers Excel --------------------------
# Causas: A,B,C  | D (sep) | E,F,G (CI) | H (sep) | I,J,K (Efectos Indirectos)
def parse_causas_xlsx(xlsx_path: str, *, sheet: Optional[str] = None, start_row: int = 3) -> Dict[str, Any]:
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet] if sheet else wb.active
    causas: Dict[str, Any] = {}
    ci_to_parent: Dict[str, str] = {}

    for row in ws.iter_rows(min_row=start_row, values_only=True):
        vals = list(row); vals += [None] * (11 - len(vals))
        A,B,C,D,E,F,G,H,I,J,K = vals[:11]

        if A:
            id_causa = str(A).strip()
            causas.setdefault(id_causa, {
                "id": id_causa,
                "descripcion": (str(B).strip() if B else None),
                "efecto_directo": {"descripcion": (str(C).strip() if C else None)},
                "causas_indirectas": {}
            })

        parent = str(E).strip() if E else None
        ci_id  = str(F).strip() if F else None
        ci_desc= str(G).strip() if G else None
        if parent and ci_id:
            base = causas.setdefault(parent, {
                "id": parent, "descripcion": None,
                "efecto_directo": {"descripcion": None},
                "causas_indirectas": {}
            })
            base["causas_indirectas"].setdefault(ci_id, {
                "id": ci_id, "descripcion": ci_desc, "efectos_indirectos": []
            })
            if ci_desc:
                base["causas_indirectas"][ci_id]["descripcion"] = ci_desc
            ci_to_parent[ci_id] = parent

            if "*" in causas:
                pend = causas["*"]["causas_indirectas"].pop(ci_id, None)
                if pend:
                    base["causas_indirectas"][ci_id]["efectos_indirectos"].extend(pend.get("efectos_indirectos", []))
                    if not base["causas_indirectas"][ci_id].get("descripcion"):
                        base["causas_indirectas"][ci_id]["descripcion"] = pend.get("descripcion")
                if not causas["*"]["causas_indirectas"]:
                    causas.pop("*", None)

        ci_ref   = str(I).strip() if I else None
        eff_id   = str(J).strip() if J else None
        eff_desc = str(K).strip() if K else None
        if ci_ref and eff_id:
            parent = ci_to_parent.get(ci_ref)
            ci_node = None
            if parent and parent in causas:
                ci_node = causas[parent]["causas_indirectas"].setdefault(
                    ci_ref, {"id":ci_ref,"descripcion":None,"efectos_indirectos":[]}
                )
            else:
                for c in causas.values():
                    if ci_ref in c["causas_indirectas"]:
                        ci_node = c["causas_indirectas"][ci_ref]; break
                if ci_node is None:
                    dummy = causas.setdefault("*", {
                        "id":"*","descripcion":None,
                        "efecto_directo":{"descripcion":None},
                        "causas_indirectas": {}
                    })
                    ci_node = dummy["causas_indirectas"].setdefault(
                        ci_ref, {"id":ci_ref,"descripcion":None,"efectos_indirectos":[]}
                    )
            ci_node["efectos_indirectos"].append({"id": eff_id, "descripcion": eff_desc})

    out: List[Dict[str, Any]] = []
    for cid, c in list(causas.items()):
        if cid == "*": continue
        c["causas_indirectas"] = list(c["causas_indirectas"].values())
        has_content = c.get("descripcion") or (c.get("efecto_directo") or {}).get("descripcion") or c["causas_indirectas"]
        if not has_content: continue
        out.append(c)

    out.sort(key=lambda x: (_num_from_id(x.get("id", "")), x.get("id", "")))
    return {"tipo": "causas", "items": out}


# Objetivos: A,B,C,D | E (sep) | F,G,H (MI) | I (sep) | J,K,L (Fines Indirectos)
def parse_objetivos_xlsx(xlsx_path: str, *, sheet: Optional[str] = None, start_row: int = 3) -> Dict[str, Any]:
    wb = load_workbook(xlsx_path, data_only=True)
    ws = wb[sheet] if sheet else wb.active
    objetivos: Dict[str, Any] = {}
    mi_to_parent: Dict[str, str] = {}

    for row in ws.iter_rows(min_row=start_row, values_only=True):
        vals = list(row); vals += [None] * (12 - len(vals))
        A,B,C,D,E,F,G,H,I,J,K,L = vals[:12]

        if A:
            id_obj = str(A).strip()
            objetivos.setdefault(id_obj, {
                "id": id_obj,
                "descripcion": (str(B).strip() if B else None),
                "medio_directo": {"descripcion": (str(C).strip() if C else None)},
                "fin_directo": {"descripcion": (str(D).strip() if D else None)},
                "medios_indirectos": {}
            })

        parent = str(F).strip() if F else None
        mi_id  = str(G).strip() if G else None
        mi_desc= str(H).strip() if H else None
        if parent and mi_id:
            base = objetivos.setdefault(parent, {
                "id": parent,
                "descripcion": None,
                "medio_directo": {"descripcion": None},
                "fin_directo": {"descripcion": None},
                "medios_indirectos": {}
            })
            base["medios_indirectos"].setdefault(mi_id, {
                "id": mi_id, "descripcion": mi_desc, "fines_indirectos": []
            })
            if mi_desc:
                base["medios_indirectos"][mi_id]["descripcion"] = mi_desc
            mi_to_parent[mi_id] = parent

            if "*" in objetivos:
                pend = objetivos["*"]["medios_indirectos"].pop(mi_id, None)
                if pend:
                    base["medios_indirectos"][mi_id]["fines_indirectos"].extend(pend.get("fines_indirectos", []))
                    if not base["medios_indirectos"][mi_id].get("descripcion"):
                        base["medios_indirectos"][mi_id]["descripcion"] = pend.get("descripcion")
                if not objetivos["*"]["medios_indirectos"]:
                    objetivos.pop("*", None)

        mi_ref  = str(J).strip() if J else None
        fi_id   = str(K).strip() if K else None
        fi_desc = str(L).strip() if L else None
        if mi_ref and fi_id:
            parent = mi_to_parent.get(mi_ref)
            mi_node = None
            if parent and parent in objetivos:
                mi_node = objetivos[parent]["medios_indirectos"].setdefault(
                    mi_ref, {"id":mi_ref,"descripcion":None,"fines_indirectos":[]}
                )
            else:
                for o in objetivos.values():
                    if mi_ref in o["medios_indirectos"]:
                        mi_node = o["medios_indirectos"][mi_ref]; break
                if mi_node is None:
                    dummy = objetivos.setdefault("*", {
                        "id":"*","descripcion":None,
                        "medio_directo":{"descripcion":None},
                        "fin_directo":{"descripcion":None},
                        "medios_indirectos": {}
                    })
                    mi_node = dummy["medios_indirectos"].setdefault(
                        mi_ref, {"id":mi_ref,"descripcion":None,"fines_indirectos":[]}
                    )
            mi_node["fines_indirectos"].append({"id": fi_id, "descripcion": fi_desc})

    out: List[Dict[str, Any]] = []
    for oid, o in list(objetivos.items()):
        if oid == "*": continue
        o["medios_indirectos"] = list(o["medios_indirectos"].values())
        has_content = o.get("descripcion") or (o.get("medio_directo") or {}).get("descripcion") or (o.get("fin_directo") or {}).get("descripcion") or o["medios_indirectos"]
        if not has_content: continue
        out.append(o)

    out.sort(key=lambda x: (_num_from_id(x.get("id", "")), x.get("id", "")))
    return {"tipo": "objetivos", "items": out}


# -------------------------- Render rápido de árboles a MD (para preview) --------------------------
def causas_tree_to_markdown(tree: Dict[str, Any]) -> str:
    if not tree or "items" not in tree: return ""
    lines = ["### Árbol de Causas y Efectos"]
    for c in tree["items"]:
        lines.append(f"- **{c['id']}**: {c.get('descripcion','') or ''}")
        ed = (c.get("efecto_directo") or {}).get("descripcion")
        if ed: lines.append(f"  - *Efecto directo:* {ed}")
        for ci in c.get("causas_indirectas", []):
            lines.append(f"  - **{ci['id']}**: {ci.get('descripcion','') or ''}")
            for ei in ci.get("efectos_indirectos", []):
                lines.append(f"    - {ei['id']}: {ei.get('descripcion','') or ''}")
    return "\n".join(lines)


def objetivos_tree_to_markdown(tree: Dict[str, Any]) -> str:
    if not tree or "items" not in tree: return ""
    lines = ["### Árbol de Objetivos, Medios y Fines"]
    for o in tree["items"]:
        lines.append(f"- **{o['id']}**: {o.get('descripcion','') or ''}")
        md = (o.get("medio_directo") or {}).get("descripcion")
        fd = (o.get("fin_directo") or {}).get("descripcion")
        if md: lines.append(f"  - *Medio directo:* {md}")
        if fd: lines.append(f"  - *Fin directo:* {fd}")
        for mi in o.get("medios_indirectos", []):
            lines.append(f"  - **{mi['id']}**: {mi.get('descripcion','') or ''}")
            for fi in mi.get("fines_indirectos", []):
                lines.append(f"    - {fi['id']}: {fi.get('descripcion','') or ''}")
    return "\n".join(lines)



def parse_mixed_sheet(filepath: str, sheet: str, start_row: int = 3) -> Dict[str, Any]:
    """
    Procesa una hoja que contiene causas y objetivos mezclados.
    Divide la hoja por bloques y usa los parsers existentes.
    """
    df = pd.read_excel(filepath, sheet_name=sheet, header=None)

    causas_df, obj_df = split_sheet_blocks(df)

    # Guardar como excels temporales para usar tus parsers existentes
    tmp_causas = "/tmp/causas_tmp.xlsx"
    tmp_obj = "/tmp/objetivos_tmp.xlsx"

    causas_df.to_excel(tmp_causas, index=False, header=False)
    obj_df.to_excel(tmp_obj, index=False, header=False)

    # Usamos tus parsers originales
    causas_tree = parse_causas_xlsx(tmp_causas, sheet=None, start_row=start_row)
    objetivos_tree = parse_objetivos_xlsx(tmp_obj, sheet=None, start_row=start_row)

    return {
        "causas": causas_tree,
        "objetivos": objetivos_tree
    }


def parse_excel_all_sheets(filepath: str, start_row: int = 3) -> Dict[str, Any]:
    wb = load_workbook(filepath, data_only=True)
    result = {}

    for sheet in wb.sheetnames:
        parsed = parse_mixed_sheet(filepath, sheet, start_row=start_row)

        # Guardar incluso si alguna parte está vacía
        result[sheet] = parsed

    return result


def process_uploaded_excel(tipo: str, filepath: str, out_dir: str) -> Dict[str, Any]:
    """
    Nuevo proceso general:
    - Ignora el parámetro 'tipo' porque ya no existen archivos separados.
    - Procesa todas las hojas.
    - Genera un JSON estructurado con causas y objetivos por hoja.
    """
    trees = parse_excel_all_sheets(filepath)

    base = os.path.splitext(os.path.basename(filepath))[0]
    out_path = save_tree_json(trees, out_dir, base)

    return {
        "json_path": out_path,
        "tree": trees,
        "preview_md": None  # opcional, podemos agregar previews por hoja si deseas
    }


# -------------------------- Conversation Flow --------------------------
conversation_flow = {
    "intro_bienvenida": {
        "prompt":
            "👋 ¡Hola!\n\n\n"
            "Soy tu asistente virtual y estoy aquí para acompañarte en la formulación de proyectos de inversión en Infraestructura de Datos (IDEC) y/o Inteligencia Artificial (IA).\n\n"
            "Te guiaré paso a paso en la formulación preliminar (borrador) del proyecto de inversión, con base en la Metodología General Ajustada (MGA) del Departamento Nacional de Planeación (DNP) y las guías de recomendaciones para la formulación de proyectos de inversión de la IDEC e IA elaboradas por la Dirección de Desarrollo Digital (DDD) del DNP, en acompañamiento la Dirección de Proyectos de Inversión(DPI) -DNP y el Ministerio TIC\n\n\n"
            "🧩 Durante el proceso:\n\n\n"
            "Te haré preguntas clave sobre tu proyecto para ayudarte a estructurarlo de manera coherente con base en los  componentes IDEC o IA que aborde tu proyecto.\n\n"
            "Esta herramienta facilitará la estructuración de los  árboles de problemas y objetivos, la definición de productos e indicadores(cadena de valor)  por componente, con la ayuda de la plantilla precargada que encontrarás en el botón \"Descargar plantillas\"  que orientará la generación de un documento borrador con la información básica del proyecto.\n\n\n"
            "📘 Recomendación:\n\n"
            "Antes o durante el uso de este asistente, revisa las guías y documentos oficiales sobre formulación de proyectos de inversión, en especial:\n\n\n"
            "El Manual de usuario del asistente que lo encuentras en el botón de \"Manual de usuario\"\n\n\n"
            "Manuales: Metodología General Ajustada para la formulación de proyectos de inversión pública en Colombia; Guía orientadora para la definición de productos: [Manuales DNP](https://www.dnp.gov.co/LaEntidad_/subdireccion-general-inversiones-seguimiento-evaluacion/direccion-proyectos-informacion-para-inversion-publica/Paginas/manuales.aspx)\n\n\n"
            "Cadena de valor: Guía de Cadena de Valor\n\n"
            "Guía para la formulación de indicadores: Guía Metodológica para la formulación de indicadores\n\n"
            "Instrumento de la MGA que consiste en la estandarización de los bienes y servicios que se pueden financiar y generar a través de los recursos públicos que son ejecutados a través de los proyectos de inversión pública. En este archivo encontrará la información estandarizada a nivel de sectores, programas y subprogramas; sectores; y productos: [Catálogo de Productos](https://colaboracion.dnp.gov.co/CDT/proyectosinformacioninversionpublica/catalogos/CATALOGO_DE_PRODUCTOS.xlsx?Web=1)\n\n\n"
            "Las guías de recomendaciones para la formulación de proyectos de inversión de la IDEC e IA (Pendiente ruta)\n\n\n"
            "Estos recursos complementan la orientación de este asistente y te ayudarán a fortalecer tu borrador de la propuesta.\n\n\n"
            "❓ Antes de continuar, ¿todo está claro? o ¿tienes algunas preguntas?",
        "options": [
            "Sí, entiendo el proceso y deseo continuar",
            "Tengo dudas respecto al proceso, me gustaría resolverlas antes de empezar"
        ],
        "next_step": "elige_vertical"
    },
    "gate_1_ciclo": {
        "prompt": "🔎 ¿Conoces el ciclo de inversión pública y las fases que lo componen?",
        "options": ["Sí, lo conozco", "No, no lo conozco"],
        "next_step": "gate_2_herramienta"
    },
    "gate_2_herramienta": {
        "prompt": "🧭 ¿Comprende que esta herramienta es de orientación y que el borrador resultante puede emplearse como insumo o apoyo en la etapa de formulación?",
        "options": ["Sí, lo comprendo", "No, no lo tengo claro"],
        "next_step": "elige_vertical"
    },

    #"rol_abierto": {
    #    "prompt": "👤 ¿Cuál es su rol dentro de la entidad (por ejemplo: Director de área, Coordinador, Profesional especializado, Analista, Asesor, Técnico operativo, Contratista de apoyo)?",
    #    "next_step": "elige_vertical"
    #},

    "elige_vertical": {
        "prompt": "💡 ¿Deseas construir un proyecto de inversión asociando componentes de tecnologías de la información y las comunicaciones en temas de Infraestructura de datos (IDEC) o Inteligencia Artificial (IA)? Puedes seleccionar una o ambas opciones.",
        "next_step": "nombre_proyecto"
    },

    "idec_componentes": {
        "prompt":
            "📚 La siguiente es la lista de los componentes que integran la IDEC, por favor selecciona los componentes que deseas incluir en tu proyecto de inversión. Selección múltiple :\n",
        "next_step": "nombre_proyecto"
    },

    "nombre_proyecto": {"prompt": "📝 ¿Cuál es el nombre del proyecto de inversión?", "next_step": "localizacion"},
    "localizacion": {"prompt": "📍 ¿Cuál es la localización en la que se enmarca el proyecto (Ejemplo: Territorial-Territorio Norte, nacional-Colombia, departamental-Cundinamarca)?", "next_step": "problema_oportunidad"},
    "problema_oportunidad": {
        "prompt": "🧩 ¿Cómo se identifica la problemática o la oportunidad a la cual se dará respuesta mediante el proyecto?\n\n"
        "**Nota:** Revisa la sección 2.1 MGA: [Documento Conceptual MGA](https://colaboracion.dnp.gov.co/CDT/proyectosinformacioninversionpublica/manuales/documento_conceptual_2023mga.pdf) donde encontrarás recomendaciones para la definición del problema central.\n\n"
        "Un proyecto nace de la intención de solucionar una situación con efectos negativos en un grupo poblacional o de aprovechar una oportunidad manifiesta dentro de un contexto particular, es decir, busca intervenir un problema para transformarlo. El foco principal de dicha problemática se denomina \"problema central\".",
        "next_step": "upload_plantilla"
    },

    "upload_plantilla": {
        "prompt": "📄 **Cargar plantilla.**\n\n"
        "1. Descargue la plantilla en la parte superior del chat.\n"
        "2. Seleccione la **PlantillaIDEC-IA.xlsx**.\n"
        "3. Diligénciela con los árboles de problemas, objetivos, productos e indicadores.\n"
        "4. Súbala en el recuadro que aparece debajo.\n\n",
        "next_step": "finalizado"
    }
}
